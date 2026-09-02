from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

from docx import Document

from app.delivery.analysis_investigation_snapshot_v1 import (
    AnalysisInvestigationSnapshotV1,
)
from app.delivery.business_document_export_v2 import (
    render_investigation_business_docx_v2,
)
from app.delivery.business_workbook_export_v2 import (
    render_investigation_business_xlsx_v2,
)
from app.delivery.investigation_report_v2 import (
    build_investigation_report_v2,
)
from app.evaluation.investigation_report_acceptance_v2 import (
    _history_item,
)


def _docx_text(payload: bytes) -> str:
    document = Document(
        BytesIO(payload)
    )

    values: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            values.append(
                paragraph.text
            )

    for table in document.tables:
        for row in table.rows:
            values.extend(
                cell.text
                for cell in row.cells
            )

    return "\n".join(values)


def _xlsx_text_and_sheets(
    payload: bytes,
) -> tuple[str, str]:
    with ZipFile(
        BytesIO(payload)
    ) as archive:
        workbook_xml = archive.read(
            "xl/workbook.xml"
        ).decode(
            "utf-8",
            errors="ignore",
        )

        parts: list[str] = []

        for name in archive.namelist():
            if (
                name == "xl/sharedStrings.xml"
                or name.startswith(
                    "xl/worksheets/"
                )
            ):
                if name.endswith(".xml"):
                    parts.append(
                        archive.read(name).decode(
                            "utf-8",
                            errors="ignore",
                        )
                    )

    return (
        workbook_xml,
        "\n".join(parts),
    )


def _report():
    return build_investigation_report_v2(
        history_item=_history_item(),
        investigation_snapshot=(
            AnalysisInvestigationSnapshotV1()
        ),
    )


def test_report_preserves_seed_comparison() -> None:
    report = _report()

    seed_view = (
        _history_item()
        .runtime_delivery_snapshot
        .console_view
    )

    assert seed_view is not None
    assert (
        report.comparison_summary
        == seed_view.comparison
    )


def test_docx_uses_business_comparison_not_seed_message() -> None:
    report = _report()

    payload = (
        render_investigation_business_docx_v2(
            report
        )
    )

    text = _docx_text(
        payload
    )

    if report.comparison_summary is not None:
        comparison = report.comparison_summary

        assert str(
            comparison.reference_value
        ) in text.replace(",", "")
        assert str(
            comparison.current_value
        ) in text.replace(",", "")
        assert (
            "Seed"
            not in text
        )
        assert (
            "Grain"
            not in text
        )


def test_workbook_has_business_core_comparison_first() -> None:
    """
    Investigation Workbook 的当前冻结顺序：

    01_决策摘要
    02_核心对比（存在 comparison 时）
    03_调查明细（存在 investigation steps 时）

    测试名称保留以避免影响现有 TESTS registry，
    但断言按最新 Business Artifact Contract 执行。
    """

    report = _report()

    payload = (
        render_investigation_business_xlsx_v2(
            report
        )
    )

    workbook_xml, xml = (
        _xlsx_text_and_sheets(
            payload
        )
    )

    assert "01_决策摘要" in workbook_xml

    if report.comparison_summary is not None:
        assert "02_核心对比" in workbook_xml

        assert (
            report.metric_definition.chinese_name
            in xml
        )
        assert (
            "环比增长额"
            in xml
            or "同比增长额"
            in xml
            or "变化额"
            in xml
        )



def test_empty_composition_sheet_is_not_created() -> None:
    report = _report()

    assert (
        report.fact_compositions
        == ()
    )

    payload = (
        render_investigation_business_xlsx_v2(
            report
        )
    )

    workbook_xml, _ = (
        _xlsx_text_and_sheets(
            payload
        )
    )

    assert (
        "构成分析"
        not in workbook_xml
    )


def test_backend_verification_sheet_is_not_created() -> None:
    report = _report()

    payload = (
        render_investigation_business_xlsx_v2(
            report
        )
    )

    workbook_xml, xml = (
        _xlsx_text_and_sheets(
            payload
        )
    )

    assert (
        'name="Verification"'
        not in workbook_xml
    )
    assert (
        "Query Plans"
        not in xml
    )
    assert (
        "Audit Events"
        not in xml
    )
    assert (
        "Evidence IDs"
        not in xml
    )


def test_business_trust_sheet_remains() -> None:
    report = _report()

    payload = (
        render_investigation_business_xlsx_v2(
            report
        )
    )

    workbook_xml, xml = (
        _xlsx_text_and_sheets(
            payload
        )
    )

    assert (
        "业务口径与可信核对"
        in workbook_xml
    )
    assert (
        report.metric_definition.definition
        in xml
    )
    assert (
        "因果结论"
        in xml
    )


TESTS = (
    test_report_preserves_seed_comparison,
    test_docx_uses_business_comparison_not_seed_message,
    test_workbook_has_business_core_comparison_first,
    test_empty_composition_sheet_is_not_created,
    test_backend_verification_sheet_is_not_created,
    test_business_trust_sheet_remains,
)


def run_acceptance() -> None:
    passed = 0
    failures: list[str] = []

    for test in TESTS:
        try:
            test()
            passed += 1
        except Exception as exc:  # noqa: BLE001
            failures.append(
                f"{test.__name__}: "
                f"{type(exc).__name__}: {exc}"
            )

    print(
        "Day94 Business Artifact UX Hardening Acceptance Summary"
    )
    print(f"Total: {len(TESTS)}")
    print(f"Passed: {passed}")
    print(f"Failed: {len(failures)}")

    for failure in failures:
        print(f"- {failure}")

    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    run_acceptance()
