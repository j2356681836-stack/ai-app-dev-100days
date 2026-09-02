from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

from docx import Document

from app.delivery.analysis_investigation_snapshot_v1 import (
    AnalysisInvestigationSnapshotV1,
)
from app.delivery.business_document_export_v2 import (
    render_investigation_business_docx_v2,
    render_periodic_business_docx_v2,
)
from app.delivery.business_workbook_export_v2 import (
    render_investigation_business_xlsx_v2,
    render_periodic_business_xlsx_v2,
)
from app.delivery.investigation_report_v2 import (
    build_investigation_report_v2,
)
from app.evaluation.investigation_report_acceptance_v2 import (
    _history_item,
)
from app.evaluation.report_export_acceptance_v2 import (
    _periodic_report,
)


def _docx_text(
    payload: bytes,
) -> str:
    document = Document(
        BytesIO(payload)
    )

    values: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            values.append(
                paragraph.text
            )

    for table in document.tables:
        for row in table.rows:
            values.extend(
                cell.text
                for cell in row.cells
            )

    return "\n".join(
        values
    )


def _xlsx_xml_text(
    payload: bytes,
) -> str:
    with ZipFile(
        BytesIO(payload)
    ) as archive:
        parts: list[str] = []

        for name in archive.namelist():
            if not name.endswith(
                ".xml"
            ):
                continue

            if not (
                name == "xl/workbook.xml"
                or name == "xl/sharedStrings.xml"
                or name.startswith(
                    "xl/worksheets/"
                )
            ):
                continue

            parts.append(
                archive.read(
                    name
                ).decode(
                    "utf-8",
                    errors="ignore",
                )
            )

    return "\n".join(
        parts
    )


def _investigation_report():
    return build_investigation_report_v2(
        history_item=_history_item(),
        investigation_snapshot=(
            AnalysisInvestigationSnapshotV1()
        ),
    )


def test_investigation_report_preserves_metric_definition() -> None:
    report = _investigation_report()

    assert (
        report.metric_definition.metric_name
        == report.metric_name
    )
    assert (
        report.metric_definition.definition
        .strip()
    )
    assert (
        report.metric_definition.formula
        .strip()
    )


def test_investigation_docx_is_business_facing() -> None:
    report = _investigation_report()

    payload = (
        render_investigation_business_docx_v2(
            report
        )
    )

    assert payload.startswith(
        b"PK"
    )

    text = _docx_text(
        payload
    )

    expected_title = {
        "fact": "AI Data Analyst · 事实分析报告",
        "composition": "AI Data Analyst · 事实分析报告",
        "comparison": "AI Data Analyst · 对比分析报告",
        "diagnostic": "AI Data Analyst · 业务调查报告",
        "investigation": "AI Data Analyst · 业务调查报告",
    }[report.analysis_mode.value]

    assert expected_title in text
    assert report.original_question in text
    assert report.metric_definition.definition in text
    assert report.metric_definition.formula in text

    assert report.history_id not in text
    assert "Query Plan" not in text
    assert "Audit Event" not in text
    assert "Contract Version" not in text
    assert "核对：reconciled" not in text

    if report.analysis_mode.value in {
        "fact",
        "composition",
    }:
        assert "核心对比" not in text
        assert "调查明细" not in text
        assert "暂不能确认" not in text

    if report.analysis_mode.value in {
        "diagnostic",
        "investigation",
    }:
        assert "决策摘要" in text




def test_investigation_xlsx_contains_business_layers_without_backend_verification() -> None:
    report = _investigation_report()

    payload = (
        render_investigation_business_xlsx_v2(
            report
        )
    )

    assert payload.startswith(
        b"PK"
    )

    xml = _xlsx_xml_text(
        payload
    )

    if report.analysis_mode.value in {
        "fact",
        "composition",
    }:
        assert "01_核心结果" in xml
    elif report.analysis_mode.value == "comparison":
        assert "01_核心对比" in xml
    else:
        assert "01_决策摘要" in xml

        if report.comparison_summary is not None:
            assert "02_核心对比" in xml
        else:
            assert "02_核心结果" in xml

    assert "业务口径与可信核对" in xml
    assert report.metric_definition.definition in xml

    assert 'name="Verification"' not in xml
    assert "Evidence IDs" not in xml
    assert "Query Plans" not in xml
    assert "Audit Events" not in xml

    if report.analysis_mode.value in {
        "fact",
        "composition",
        "comparison",
    }:
        assert "因果结论" not in xml




def test_periodic_docx_is_business_facing() -> None:
    report = _periodic_report()

    payload = (
        render_periodic_business_docx_v2(
            report
        )
    )

    assert payload.startswith(
        b"PK"
    )

    text = _docx_text(
        payload
    )

    for token in (
        "AI Data Analyst",
        "本期经营摘要",
        "经营观察",
        "经营概览",
        report.metrics[0].spec.chinese_name,
        report.metrics[0].spec.purpose,
    ):
        assert token in text

    assert report.contract_version not in text



def test_periodic_xlsx_contains_business_kpi_and_trust_layers() -> None:
    report = _periodic_report()

    payload = (
        render_periodic_business_xlsx_v2(
            report
        )
    )

    assert payload.startswith(
        b"PK"
    )

    xml = _xlsx_xml_text(
        payload
    )

    for token in (
        "01_经营摘要",
        "02_经营概览",
        "03_销售驱动",
        "04_客户健康",
        "经营观察",
        "报表周期",
        "锚点日期",
        "参考期",
        "当前期",
        report.metrics[0].spec.chinese_name,
        "业务口径与可信核对",
    ):
        assert token in xml

    assert "01_KPI" not in xml

    if report.driver_reconciliations:
        assert "05_驱动核对" in xml

    assert (
        report.metrics[0].current_evidence_id
        not in xml
    )
    assert (
        report.metrics[0].reference_evidence_id
        not in xml
    )


def test_business_export_is_repeatable_at_semantic_level() -> None:
    report = _investigation_report()

    first_doc = _docx_text(
        render_investigation_business_docx_v2(
            report
        )
    )
    second_doc = _docx_text(
        render_investigation_business_docx_v2(
            report
        )
    )

    assert (
        first_doc
        == second_doc
    )

    first_xlsx = _xlsx_xml_text(
        render_investigation_business_xlsx_v2(
            report
        )
    )
    second_xlsx = _xlsx_xml_text(
        render_investigation_business_xlsx_v2(
            report
        )
    )

    assert (
        first_xlsx
        == second_xlsx
    )


TESTS = (
    test_investigation_report_preserves_metric_definition,
    test_investigation_docx_is_business_facing,
    test_investigation_xlsx_contains_business_layers_without_backend_verification,
    test_periodic_docx_is_business_facing,
    test_periodic_xlsx_contains_business_kpi_and_trust_layers,
    test_business_export_is_repeatable_at_semantic_level,
)


def run_acceptance() -> None:
    passed = 0
    failures: list[str] = []

    for test in TESTS:
        try:
            test()
            passed += 1
        except Exception as exc:  # noqa: BLE001
            failures.append(
                f"{test.__name__}: "
                f"{type(exc).__name__}: {exc}"
            )

    print(
        "Day94 Business Artifact Export Acceptance Summary"
    )
    print(
        f"Total: {len(TESTS)}"
    )
    print(
        f"Passed: {passed}"
    )
    print(
        f"Failed: {len(failures)}"
    )

    for failure in failures:
        print(
            f"- {failure}"
        )

    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    run_acceptance()
