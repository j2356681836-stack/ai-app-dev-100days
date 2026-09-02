from __future__ import annotations

import inspect
from datetime import date
from decimal import Decimal
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from app.delivery.analysis_investigation_snapshot_v1 import (
    AnalysisInvestigationSnapshotV1,
)
from app.delivery.business_document_export_v2 import (
    render_investigation_business_docx_v2,
)
from app.delivery.business_workbook_export_v2 import (
    render_investigation_business_xlsx_v2,
)
from app.delivery.fact_composition_delivery_v2 import (
    FactCompositionDimensionV2,
    build_fact_composition_projection_v2,
)
from app.delivery.investigation_report_v2 import (
    build_investigation_report_v2,
)
from app.evaluation.investigation_report_acceptance_v2 import (
    _history_item,
)
from app.semantic_layer.analysis_mode_contract_v2 import (
    AnalysisModeV2,
)
from app.semantic_layer.time_comparison_contract_v2 import (
    TimeWindowReferenceV2,
)
from app.ui import decision_console_app as app


WINDOW = TimeWindowReferenceV2(
    start_date=date(2025, 1, 1),
    end_date=date(2025, 12, 31),
)


def _composition(
    dimension: FactCompositionDimensionV2,
):
    if dimension == FactCompositionDimensionV2.PEOPLE:
        rows = (
            {
                "membership_segment": "Gold",
                "gmv": Decimal("600"),
            },
            {
                "membership_segment": "NON_MEMBER",
                "gmv": Decimal("400"),
            },
        )
        plan_name = "gmv_membership_level_v2"
    elif dimension == FactCompositionDimensionV2.CHANNEL:
        rows = (
            {
                "channel_name": "渠道A",
                "gmv": Decimal("600"),
            },
            {
                "channel_name": "渠道B",
                "gmv": Decimal("400"),
            },
        )
        plan_name = "gmv_channel_v2"
    else:
        rows = (
            {
                "category": "护肤",
                "gmv": Decimal("700"),
            },
            {
                "category": "彩妆",
                "gmv": Decimal("300"),
            },
        )
        plan_name = "gmv_category_v2"

    return build_fact_composition_projection_v2(
        dimension=dimension,
        metric_name="gmv",
        overall_value=Decimal("1000"),
        analysis_window=WINDOW,
        scope_summary="test",
        rows=rows,
        evidence_id=f"ev-{dimension.value}",
        plan_name=plan_name,
        audit_event_id=f"audit-{dimension.value}",
    )


def _report():
    base = build_investigation_report_v2(
        history_item=_history_item(),
        investigation_snapshot=(
            AnalysisInvestigationSnapshotV1()
        ),
    )

    return base.model_copy(
        update={
            "analysis_mode": AnalysisModeV2.FACT,
            "comparison_summary": None,
            "fact_compositions": (
                _composition(
                    FactCompositionDimensionV2.PEOPLE
                ),
                _composition(
                    FactCompositionDimensionV2.CHANNEL
                ),
                _composition(
                    FactCompositionDimensionV2.CATEGORY
                ),
            ),
        }
    )


def _docx_text(payload: bytes) -> str:
    document = Document(
        BytesIO(payload)
    )
    values: list[str] = []

    for paragraph in document.paragraphs:
        values.append(
            paragraph.text
        )

    for table in document.tables:
        for row in table.rows:
            values.extend(
                cell.text
                for cell in row.cells
            )

    return "\n".join(values)


def _xlsx_xml(payload: bytes) -> str:
    with ZipFile(
        BytesIO(payload)
    ) as archive:
        return "\n".join(
            archive.read(name).decode(
                "utf-8",
                errors="ignore",
            )
            for name in archive.namelist()
            if (
                name.endswith(".xml")
                and (
                    name == "xl/workbook.xml"
                    or name == "xl/sharedStrings.xml"
                    or name.startswith(
                        "xl/worksheets/"
                    )
                )
            )
        )


def test_business_ui_does_not_release_people_composition() -> None:
    source = inspect.getsource(
        app._render_fact_composition_section_v2
    )

    assert (
        "FactCompositionDimensionV2.PEOPLE"
        in source
    )
    assert "人群构成仍在口径校准中" in source


def test_business_labels_no_longer_use_people_goods_place_copy() -> None:
    source = inspect.getsource(
        app._composition_dimension_label_v2
    )

    assert '"渠道构成"' in source
    assert '"品类构成"' in source
    assert '"人群构成"' in source

    for token in (
        '"人｜',
        '"货｜',
        '"场｜',
    ):
        assert token not in source


def test_cross_check_filters_deferred_people_dimension() -> None:
    source = inspect.getsource(
        app._render_fact_composition_cross_check_v2
    )

    assert (
        "FactCompositionDimensionV2.PEOPLE"
        in source
    )


def test_fact_verification_filters_deferred_people_dimension() -> None:
    source = inspect.getsource(
        app._render_business_verification_dialog_v2
    )

    assert (
        "FactCompositionDimensionV2.PEOPLE"
        in source
    )


def test_docx_has_public_blocks_and_trusted_summary_rows() -> None:
    text = _docx_text(
        render_investigation_business_docx_v2(
            _report()
        )
    )

    assert "渠道构成" in text
    assert "品类构成" in text
    assert "人群构成" not in text
    assert "会员构成" not in text

    # Channel + Category each contribute one trusted summary row.
    assert text.count("汇总") >= 2


def test_xlsx_has_separate_public_blocks_and_no_people_block() -> None:
    xml = _xlsx_xml(
        render_investigation_business_xlsx_v2(
            _report()
        )
    )

    assert "渠道构成" in xml
    assert "品类构成" in xml
    assert "人群构成" not in xml
    assert "membership_level" not in xml
    assert "构成分析" in xml
    assert "汇总" in xml


TESTS = (
    test_business_ui_does_not_release_people_composition,
    test_business_labels_no_longer_use_people_goods_place_copy,
    test_cross_check_filters_deferred_people_dimension,
    test_fact_verification_filters_deferred_people_dimension,
    test_docx_has_public_blocks_and_trusted_summary_rows,
    test_xlsx_has_separate_public_blocks_and_no_people_block,
)


def run_acceptance() -> None:
    passed = 0
    failures: list[str] = []

    for test in TESTS:
        try:
            test()
            passed += 1
        except Exception as exc:  # noqa: BLE001
            failures.append(
                f"{test.__name__}: "
                f"{type(exc).__name__}: {exc}"
            )

    print(
        "Day94 Fact Composition Public Freeze Acceptance"
    )
    print(f"Total: {len(TESTS)}")
    print(f"Passed: {passed}")
    print(f"Failed: {len(failures)}")

    for failure in failures:
        print(f"- {failure}")

    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    run_acceptance()
