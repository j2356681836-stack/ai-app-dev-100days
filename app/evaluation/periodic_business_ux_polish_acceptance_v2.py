from __future__ import annotations

import inspect
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from app.delivery.business_document_export_v2 import (
    _add_periodic_metric_table,
    render_periodic_business_docx_v2,
)
from app.delivery.business_workbook_export_v2 import (
    render_periodic_business_xlsx_v2,
)
from app.evaluation.report_export_acceptance_v2 import (
    _periodic_report,
)
from app.ui import decision_console_app as ui


def _docx_text(payload: bytes) -> str:
    document = Document(BytesIO(payload))
    values: list[str] = []

    for paragraph in document.paragraphs:
        values.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            values.extend(
                cell.text
                for cell in row.cells
            )

    return "\n".join(values)


def _xlsx_xml(payload: bytes) -> str:
    with ZipFile(BytesIO(payload)) as archive:
        parts: list[str] = []

        for name in archive.namelist():
            if (
                name.endswith(".xml")
                and (
                    name == "xl/workbook.xml"
                    or name == "xl/sharedStrings.xml"
                    or name.startswith(
                        "xl/worksheets/"
                    )
                )
            ):
                parts.append(
                    archive.read(name).decode(
                        "utf-8",
                        errors="ignore",
                    )
                )

    return "\n".join(parts)


def test_periodic_anchor_is_shared_across_cadence() -> None:
    daily = ui._periodic_anchor_state_key_v2("daily")
    weekly = ui._periodic_anchor_state_key_v2("weekly")
    monthly = ui._periodic_anchor_state_key_v2("monthly")

    assert daily == weekly == monthly
    assert daily == "periodic_anchor_date::shared"


def test_periodic_word_does_not_list_release_status_column() -> None:
    source = inspect.getsource(
        _add_periodic_metric_table
    )

    assert "cols=5" in source
    assert '"状态"' not in source
    assert '"可释放"' not in source


def test_periodic_word_is_business_facing() -> None:
    report = _periodic_report()
    payload = render_periodic_business_docx_v2(
        report
    )
    text = _docx_text(payload)

    assert "报表状态：" not in text
    assert "partial_ready" not in text
    assert "PARTIAL_READY" not in text
    assert "reconciled" not in text

    if report.driver_reconciliations:
        assert "已对账" in text


def test_periodic_excel_names_cadence_and_time_windows() -> None:
    report = _periodic_report()
    payload = render_periodic_business_xlsx_v2(
        report
    )
    xml = _xlsx_xml(payload)

    cadence_label = {
        "daily": "日报",
        "weekly": "周报",
        "monthly": "月报",
    }[report.cadence.value]

    reference = report.comparison.reference_window
    current = report.comparison.current_window

    assert "报表周期" in xml
    assert cadence_label in xml
    assert "锚点日期" in xml
    assert "参考期" in xml
    assert "当前期" in xml
    assert str(reference.start_date) in xml
    assert str(reference.end_date) in xml
    assert str(current.start_date) in xml
    assert str(current.end_date) in xml


TESTS = (
    test_periodic_anchor_is_shared_across_cadence,
    test_periodic_word_does_not_list_release_status_column,
    test_periodic_word_is_business_facing,
    test_periodic_excel_names_cadence_and_time_windows,
)


def run_acceptance() -> None:
    passed = 0
    failures: list[str] = []

    for test in TESTS:
        try:
            test()
            passed += 1
        except Exception as exc:  # noqa: BLE001
            failures.append(
                f"{test.__name__}: "
                f"{type(exc).__name__}: {exc}"
            )

    print(
        "Day94 Periodic Business UX Polish Acceptance"
    )
    print(f"Total: {len(TESTS)}")
    print(f"Passed: {passed}")
    print(f"Failed: {len(failures)}")

    for failure in failures:
        print(f"- {failure}")

    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    run_acceptance()
