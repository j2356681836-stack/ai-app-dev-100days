from __future__ import annotations

import inspect
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from app.delivery.analysis_investigation_snapshot_v1 import (
    AnalysisInvestigationSnapshotV1,
)
from app.delivery.business_document_export_v2 import (
    render_investigation_business_docx_v2,
)
from app.delivery.business_workbook_export_v2 import (
    render_investigation_business_xlsx_v2,
)
from app.delivery.investigation_report_v2 import (
    build_investigation_report_v2,
)
from app.evaluation.investigation_report_acceptance_v2 import (
    _history_item,
)
from app.semantic_layer.analysis_mode_contract_v2 import (
    AnalysisModeV2,
)
from app.ui import decision_console_app as app


def _report():
    return build_investigation_report_v2(
        history_item=_history_item(),
        investigation_snapshot=(
            AnalysisInvestigationSnapshotV1()
        ),
    )


def _docx_text(payload: bytes) -> str:
    document = Document(
        BytesIO(payload)
    )
    values: list[str] = []

    for paragraph in document.paragraphs:
        values.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            values.extend(
                cell.text
                for cell in row.cells
            )

    return "\n".join(values)


def _xlsx_xml(payload: bytes) -> str:
    with ZipFile(BytesIO(payload)) as archive:
        parts: list[str] = []

        for name in archive.namelist():
            if (
                name.endswith(".xml")
                and (
                    name == "xl/workbook.xml"
                    or name == "xl/sharedStrings.xml"
                    or name.startswith(
                        "xl/worksheets/"
                    )
                )
            ):
                parts.append(
                    archive.read(name).decode(
                        "utf-8",
                        errors="ignore",
                    )
                )

    return "\n".join(parts)


def test_report_preserves_requested_analysis_mode() -> None:
    history = _history_item()
    report = build_investigation_report_v2(
        history_item=history,
        investigation_snapshot=(
            AnalysisInvestigationSnapshotV1()
        ),
    )

    assert (
        report.analysis_mode
        == history.runtime_delivery_snapshot.requested_analysis_mode
    )


def test_fact_export_has_no_comparison_or_investigation_semantics() -> None:
    report = _report().model_copy(
        update={
            "analysis_mode": AnalysisModeV2.FACT,
            "comparison_summary": None,
            "investigation_steps": (),
            "user_exploration_steps": (),
        }
    )

    text = _docx_text(
        render_investigation_business_docx_v2(
            report
        )
    )
    xml = _xlsx_xml(
        render_investigation_business_xlsx_v2(
            report
        )
    )

    assert "事实分析报告" in text
    assert "核心对比" not in text
    assert "调查结果" not in text
    assert "暂不能确认" not in text
    assert "下一步建议" not in text
    assert "01_核心结果" in xml
    assert "01_核心对比" not in xml
    assert "因果结论" not in xml


def test_comparison_export_has_no_investigation_sections() -> None:
    report = _report().model_copy(
        update={
            "analysis_mode": AnalysisModeV2.COMPARISON,
            "investigation_steps": (),
            "user_exploration_steps": (),
        }
    )

    text = _docx_text(
        render_investigation_business_docx_v2(
            report
        )
    )
    xml = _xlsx_xml(
        render_investigation_business_xlsx_v2(
            report
        )
    )

    assert "对比分析报告" in text
    assert "调查结果" not in text
    assert "暂不能确认" not in text
    assert "因果结论" not in xml


def test_investigation_export_keeps_investigation_profile() -> None:
    report = _report().model_copy(
        update={
            "analysis_mode": AnalysisModeV2.INVESTIGATION,
        }
    )

    text = _docx_text(
        render_investigation_business_docx_v2(
            report
        )
    )
    xml = _xlsx_xml(
        render_investigation_business_xlsx_v2(
            report
        )
    )

    assert "业务调查报告" in text
    assert "业务口径与可信边界" in text
    assert "因果结论" in xml


def test_fact_verification_uses_dialog_not_expander() -> None:
    source = inspect.getsource(
        app._render_fact_verification_v2
    )

    assert "st.button(" in source
    assert "st.expander(" not in source
    assert (
        "_render_business_verification_dialog_v2"
        in source
    )


def test_business_verification_dialog_has_no_engineering_ids() -> None:
    source = inspect.getsource(
        app._render_business_verification_dialog_v2
    )

    # 只禁止真正可能渲染到业务 Dialog 的工程字段。
    # Docstring / internal comments 可以描述“工程信息留在工程视图”，
    # 不能因此被误判为业务 UI 泄漏。
    forbidden_ui_patterns = (
        'st.write("Evidence ID',
        'st.write("Query Plan',
        'st.write("Audit Event',
        'st.markdown("**Governed Evidence',
        "released_field_names",
        "released_row_count",
    )

    for pattern in forbidden_ui_patterns:
        assert pattern not in source

    for target in (
        '"fact"',
        '"comparison"',
        '"periodic"',
        '"periodic_r12"',
        '"periodic_contribution"',
        '"step::"',
    ):
        assert target in source



def test_periodic_verification_uses_dialog() -> None:
    report_source = inspect.getsource(
        app._render_periodic_business_report_v2
    )
    comparison_source = inspect.getsource(
        app._render_periodic_comparison_business
    )

    assert "验证 R12 客户指标" in report_source
    assert "验证这份周期报表" in report_source
    assert '"periodic_r12"' in report_source
    assert '"periodic"' in report_source
    assert "Governed Evidence IDs" not in report_source

    assert "Audit Event" not in comparison_source
    assert '"periodic_contribution"' in comparison_source


def test_engineering_view_owns_internal_entry_and_lineage() -> None:
    source = inspect.getsource(
        app._render_engineering_view
    )
    agentic_source = inspect.getsource(
        app._render_agentic_business_section
    )

    assert "入口类型：" in source
    assert "查看内部入口合同 Payload" in source
    assert (
        "_render_active_analysis_evidence_lineage_v1()"
        in source
    )
    assert (
        "_render_active_analysis_evidence_lineage_v1()"
        not in agentic_source
    )


def test_export_ui_is_mode_aware() -> None:
    source = inspect.getsource(
        app._render_investigation_report_export_v2
    )

    assert "report_payload.analysis_mode" in source
    assert "当前导出类型" in source
    assert "InvestigationReportV2；" not in source


TESTS = (
    test_report_preserves_requested_analysis_mode,
    test_fact_export_has_no_comparison_or_investigation_semantics,
    test_comparison_export_has_no_investigation_sections,
    test_investigation_export_keeps_investigation_profile,
    test_fact_verification_uses_dialog_not_expander,
    test_business_verification_dialog_has_no_engineering_ids,
    test_periodic_verification_uses_dialog,
    test_engineering_view_owns_internal_entry_and_lineage,
    test_export_ui_is_mode_aware,
)


def run_acceptance() -> None:
    passed = 0
    failures: list[str] = []

    for test in TESTS:
        try:
            test()
            passed += 1
        except Exception as exc:  # noqa: BLE001
            failures.append(
                f"{test.__name__}: "
                f"{type(exc).__name__}: {exc}"
            )

    print(
        "Day94 Presentation & Verification Consistency Acceptance"
    )
    print(f"Total: {len(TESTS)}")
    print(f"Passed: {passed}")
    print(f"Failed: {len(failures)}")

    for failure in failures:
        print(f"- {failure}")

    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    run_acceptance()
