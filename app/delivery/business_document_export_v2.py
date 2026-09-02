from __future__ import annotations

from io import BytesIO
from decimal import Decimal

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.delivery.investigation_report_v2 import (
    InvestigationReportV2,
)
from app.semantic_layer.analysis_mode_contract_v2 import (
    AnalysisModeV2,
)
from app.delivery.periodic_business_report_v2 import (
    PeriodicBusinessReportV2,
    PeriodicMetricDisplayKindV2,
    PeriodicMetricSectionV2,
    PeriodicMetricStatusV2,
)


BUSINESS_DOCUMENT_EXPORT_VERSION = "business_document_export_v2_1"


def _decimal_text(value: Decimal | None) -> str:
    if value is None:
        return "—"
    return f"{value:,.2f}"


def _ratio_text(value: Decimal | None) -> str:
    if value is None:
        return "—"
    return f"{value * Decimal('100'):.2f}%"


def _comparison_change_labels(comparison_type) -> tuple[str, str]:
    value = comparison_type.value

    if value == "mom":
        return "环比增长额", "环比增长率"
    if value == "yoy":
        return "同比增长额", "同比增长率"
    if value == "wow":
        return "周环比增长额", "周环比增长率"
    if value == "dod":
        return "日环比增长额", "日环比增长率"

    return "变化额", "变化率"


def _period_label(window) -> str:
    start = window.start_date
    end = window.end_date

    if (
        start.year == end.year
        and start.month == end.month
        and start.day == 1
        and end.day >= 28
    ):
        return f"{start.year}年{start.month}月"

    if start == end:
        return str(start)

    return f"{start} 至 {end}"


def _periodic_metric_value_text(
    snapshot,
    *,
    reference: bool,
) -> str:
    if snapshot.status != PeriodicMetricStatusV2.READY:
        return "不可释放"

    value = (
        snapshot.reference_value
        if reference
        else snapshot.current_value
    )

    if snapshot.spec.display_kind == PeriodicMetricDisplayKindV2.RATIO:
        return _ratio_text(value)

    if snapshot.spec.display_kind == PeriodicMetricDisplayKindV2.COUNT:
        if value is None:
            return "—"
        return f"{value:,.0f}"

    return _decimal_text(value)


def _periodic_change_text(snapshot) -> str:
    if snapshot.status != PeriodicMetricStatusV2.READY:
        return "—"

    if snapshot.spec.display_kind == PeriodicMetricDisplayKindV2.RATIO:
        value = snapshot.percentage_point_change
        if value is None:
            return "—"
        return f"{value:+.2f} 个百分点"

    if snapshot.absolute_change is None:
        return "—"

    if snapshot.spec.display_kind == PeriodicMetricDisplayKindV2.COUNT:
        return f"{snapshot.absolute_change:+,.0f}"

    return f"{snapshot.absolute_change:+,.2f}"


def _set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Microsoft YaHei",
    )

    for style_name, size in (
        ("Title", 20),
        ("Heading 1", 15),
        ("Heading 2", 12),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Microsoft YaHei",
        )


def _add_title(
    document: Document,
    title: str,
    subtitle: str | None = None,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(20)

    if subtitle:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(subtitle)
        run.font.size = Pt(10)


def _add_bullets(
    document: Document,
    items: tuple[str, ...] | list[str],
) -> None:
    if not items:
        document.add_paragraph("—")
        return

    for item in items:
        document.add_paragraph(
            str(item),
            style="List Bullet",
        )


def _set_table_header(row) -> None:
    for cell in row.cells:
        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _add_comparison_section(
    document: Document,
    report: InvestigationReportV2,
) -> None:
    comparison = report.comparison_summary

    if comparison is None:
        document.add_paragraph(
            "本次交付未包含独立的整体时间比较结果。"
        )
        return

    change_label, rate_label = (
        _comparison_change_labels(
            comparison.comparison_type
        )
    )

    table = document.add_table(
        rows=1,
        cols=5,
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = (
        "指标",
        _period_label(comparison.reference_window),
        _period_label(comparison.current_window),
        change_label,
        rate_label,
    )

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    _set_table_header(table.rows[0])

    row = table.add_row().cells
    row[0].text = report.metric_definition.chinese_name
    row[1].text = _decimal_text(comparison.reference_value)
    row[2].text = _decimal_text(comparison.current_value)
    row[3].text = _decimal_text(comparison.absolute_change)
    row[4].text = _ratio_text(comparison.relative_change)

    document.add_paragraph(
        f"{_period_label(comparison.current_window)}"
        f"{report.metric_definition.chinese_name}较"
        f"{_period_label(comparison.reference_window)}"
        f"变化 {_decimal_text(comparison.absolute_change)}，"
        f"{rate_label} {_ratio_text(comparison.relative_change)}。"
    )


def _dimension_label(value: str) -> str:
    return {
        "channel": "渠道",
        "category": "品类",
        "region": "城市",
        "area": "大区",
        "province": "省级地区",
        "city": "城市",
        "campaign": "活动实例",
    }.get(value, value)


def _add_change_step(
    document: Document,
    *,
    step,
    sequence: int,
    source_label: str,
) -> None:
    """
    Investigation detail block with trusted summary row.

    汇总行直接读取 result 已冻结的 focus values / focus_delta；
    不在导出层重新聚合成员。
    """

    result = step.result
    label = _dimension_label(
        result.dimension_name.value
    )

    document.add_heading(
        f"{sequence}. {label}变化",
        level=2,
    )
    document.add_paragraph(
        f"方向来源：{source_label}"
    )

    if step.assessment is not None:
        document.add_paragraph(
            f"本轮结论：{step.assessment.conclusion}"
        )

    table = document.add_table(
        rows=1,
        cols=5,
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = (
        label,
        "参考期",
        "当前期",
        "变化额",
        "占本轮变化",
    )
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    _set_table_header(table.rows[0])

    for member in result.members:
        row = table.add_row().cells
        row[0].text = member.member_label
        row[1].text = _decimal_text(
            member.reference_value
        )
        row[2].text = _decimal_text(
            member.current_value
        )
        row[3].text = _decimal_text(
            member.delta
        )
        row[4].text = _ratio_text(
            member.share_of_focus_delta
        )

    summary = table.add_row().cells
    summary[0].text = "汇总"
    summary[1].text = _decimal_text(
        result.reference_focus_value
    )
    summary[2].text = _decimal_text(
        result.current_focus_value
    )
    summary[3].text = _decimal_text(
        result.focus_delta
    )
    summary[4].text = (
        "100.00%"
        if (
            result.focus_delta != 0
            and result.reconciliation_status.value
            == "reconciled"
        )
        else "—"
    )

    for cell in summary:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    status = {
        "reconciled": "已对账",
        "not_reconciled": "未完全对账",
    }.get(
        result.reconciliation_status.value,
        result.reconciliation_status.value,
    )

    document.add_paragraph(
        f"核对状态：{status}；"
        f"成员变化合计 {_decimal_text(result.sum_member_delta)}；"
        f"未解释差额 {_decimal_text(result.unexplained_remainder)}。"
    )




_FACT_COUNT_METRICS_V2 = frozenset(
    {
        "order_count",
        "buyer_count",
        "units_sold",
        "repeat_customer_count",
        "multi_order_customer_count",
        "brand_paid_new_customer_count",
        "channel_paid_new_customer_count",
        "r12_base_customer_count",
        "r12_repurchase_customer_count",
    }
)


def _report_mode_group_v2(
    report: InvestigationReportV2,
) -> str:
    if report.analysis_mode in {
        AnalysisModeV2.FACT,
        AnalysisModeV2.COMPOSITION,
    }:
        return "fact"

    if (
        report.analysis_mode
        == AnalysisModeV2.COMPARISON
    ):
        return "comparison"

    return "investigation"


def _business_report_title_v2(
    report: InvestigationReportV2,
) -> str:
    return {
        "fact": "AI Data Analyst · 事实分析报告",
        "comparison": "AI Data Analyst · 对比分析报告",
        "investigation": "AI Data Analyst · 业务调查报告",
    }[_report_mode_group_v2(report)]


def _evidence_sufficiency_label_v2(
    value,
) -> str:
    raw = (
        value.value
        if hasattr(value, "value")
        else str(value)
    )

    return {
        "sufficient_for_current_scope": "当前范围证据充分",
        "partial": "部分充分",
        "insufficient": "证据不足",
    }.get(
        raw,
        raw,
    )


def _fact_metric_text_v2(
    metric_name: str,
    value: Decimal | None,
) -> str:
    if value is None:
        return "—"

    if (
        metric_name.strip().lower()
        in _FACT_COUNT_METRICS_V2
    ):
        return f"{value:,.0f}"

    return _decimal_text(value)



def _is_public_fact_composition_v2(
    result,
) -> bool:
    """
    Business Artifact public-release filter.

    Legacy PEOPLE snapshots remain available to technical contracts,
    but are not exported until the frozen lifecycle + pre-window
    membership semantics are implemented.
    """

    return result.dimension.value != "membership_level"


def _public_fact_compositions_v2(
    report: InvestigationReportV2,
):
    return tuple(
        result
        for result in report.fact_compositions
        if _is_public_fact_composition_v2(
            result
        )
    )

def _fact_dimension_label_v2(
    value: str,
) -> str:
    return {
        "membership_level": "人群构成",
        "channel": "渠道构成",
        "category": "品类构成",
        "region": "地区构成",
    }.get(
        value,
        value,
    )



def _add_fact_composition_sections_v2(
    document: Document,
    report: InvestigationReportV2,
) -> None:
    """
    Business DOCX Composition。

    每个维度独立成表；
    最后一行“汇总”直接使用 Delivery trusted Overall，
    不在导出层对成员重新求和。
    """

    for result in _public_fact_compositions_v2(
        report
    ):
        document.add_heading(
            _fact_dimension_label_v2(
                result.dimension.value
            ),
            level=2,
        )

        table = document.add_table(
            rows=1,
            cols=3,
        )
        table.style = "Table Grid"
        table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )

        headers = (
            "成员",
            report.metric_definition.chinese_name,
            "构成占比",
        )

        for index, header in enumerate(
            headers
        ):
            table.rows[0].cells[index].text = (
                header
            )
        _set_table_header(
            table.rows[0]
        )

        for member in result.members:
            row = table.add_row().cells
            row[0].text = member.member_label
            row[1].text = _fact_metric_text_v2(
                result.metric_name,
                member.value,
            )
            row[2].text = _ratio_text(
                member.share
            )

        summary = table.add_row().cells
        summary[0].text = "汇总"
        summary[1].text = _fact_metric_text_v2(
            result.metric_name,
            result.overall_value,
        )
        summary[2].text = (
            "100.00%"
            if (
                result.overall_value
                is not None
                and result.overall_value != 0
            )
            else "未定义"
        )

        for cell in summary:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        status = (
            "已对账"
            if (
                result.reconciliation_status.value
                == "reconciled"
            )
            else "未完全对账"
        )

        document.add_paragraph(
            f"构成合计："
            f"{_fact_metric_text_v2(result.metric_name, result.member_sum)}；"
            f"可信 Overall："
            f"{_fact_metric_text_v2(result.metric_name, result.overall_value)}；"
            f"未解释差额："
            f"{_fact_metric_text_v2(result.metric_name, result.unexplained_remainder)}；"
            f"核对状态：{status}。"
        )



def _add_business_trust_section_v2(
    document: Document,
    report: InvestigationReportV2,
) -> None:
    metric = report.metric_definition

    table = document.add_table(
        rows=0,
        cols=2,
    )
    table.style = "Table Grid"
    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    rows = (
        (
            "指标",
            f"{metric.chinese_name}（{metric.metric_name}）",
        ),
        ("定义", metric.definition),
        ("公式", metric.formula),
        ("基础粒度", metric.grain),
        (
            "过滤条件",
            (
                "；".join(metric.filters)
                if metric.filters
                else "无额外过滤条件"
            ),
        ),
        (
            "证据充分性",
            _evidence_sufficiency_label_v2(
                report.executive_brief.evidence_sufficiency
            ),
        ),
    )

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
        cells[0].paragraphs[0].runs[0].bold = True

    limitations = tuple(
        item.detail
        for item in report.executive_brief.limitations
    )

    if limitations:
        document.add_paragraph(
            "必要限制："
        )
        _add_bullets(
            document,
            list(limitations),
        )

    if (
        _report_mode_group_v2(report)
        == "investigation"
    ):
        document.add_paragraph(
            "说明：数值分解用于定位变化主要发生在哪里；"
            "除非存在额外反事实或实验类证据，"
            "不能据此直接宣称业务因果。"
        )


def _dedupe_business_text_v2(
    values,
) -> tuple[str, ...]:
    result: list[str] = []
    seen: set[str] = set()

    for value in values:
        text = str(value).strip()
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)

    return tuple(result)


def _investigation_all_steps_v2(
    report: InvestigationReportV2,
):
    return (
        *report.investigation_steps,
        *report.user_exploration_steps,
    )


def _investigation_confirmed_v2(
    report: InvestigationReportV2,
) -> tuple[str, ...]:
    """
    Investigation“可以确认”只保留调查步骤新增确认项。

    Overall Comparison 已在“核心对比”出现，不在这里重复。
    """

    values: list[str] = []

    for step in _investigation_all_steps_v2(
        report
    ):
        if step.assessment is None:
            continue
        values.extend(
            step.assessment.can_confirm
        )

    return _dedupe_business_text_v2(
        values
    )


def _investigation_cannot_confirm_v2(
    report: InvestigationReportV2,
) -> tuple[str, ...]:
    values: list[str] = [
        *report.executive_brief.candidate_hypotheses,
        *report.executive_brief.unknowns,
    ]

    for step in _investigation_all_steps_v2(
        report
    ):
        if step.assessment is None:
            continue
        values.extend(
            step.assessment.cannot_confirm
        )

    return _dedupe_business_text_v2(
        values
    )


def _investigation_next_steps_v2(
    report: InvestigationReportV2,
) -> tuple[str, ...]:
    values: list[str] = list(
        report.executive_brief.recommended_checks
    )

    for step in _investigation_all_steps_v2(
        report
    ):
        if step.assessment is None:
            continue
        values.append(
            step.assessment.next_step_recommendation
        )

    return _dedupe_business_text_v2(
        values
    )


def _add_investigation_decision_summary_v2(
    document: Document,
    report: InvestigationReportV2,
) -> None:
    """
    只投影现有 Comparison / Assessment / Sufficiency。

    不调用 LLM，不重新生成原因，不重新计算 KPI。
    """

    table = document.add_table(
        rows=0,
        cols=2,
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    comparison = report.comparison_summary

    if comparison is not None:
        _, rate_label = _comparison_change_labels(
            comparison.comparison_type
        )
        cells = table.add_row().cells
        cells[0].text = "整体变化"
        cells[1].text = (
            f"{_period_label(comparison.current_window)}较"
            f"{_period_label(comparison.reference_window)}"
            f"变化 {_decimal_text(comparison.absolute_change)}；"
            f"{rate_label} {_ratio_text(comparison.relative_change)}。"
        )

    cells = table.add_row().cells
    cells[0].text = "证据状态"
    cells[1].text = _evidence_sufficiency_label_v2(
        report.executive_brief.evidence_sufficiency
    )

    finding_index = 0
    for step in _investigation_all_steps_v2(
        report
    ):
        if (
            step.assessment is None
            or not step.assessment.conclusion.strip()
        ):
            continue

        finding_index += 1
        cells = table.add_row().cells
        cells[0].text = f"调查发现 {finding_index}"
        cells[1].text = step.assessment.conclusion

    next_steps = _investigation_next_steps_v2(
        report
    )

    if next_steps:
        cells = table.add_row().cells
        cells[0].text = "优先下一步"
        cells[1].text = next_steps[0]

    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True


def _periodic_metric_lookup_v2(
    report: PeriodicBusinessReportV2,
) -> dict[str, object]:
    return {
        item.spec.metric_name: item
        for item in report.metrics
    }


def _periodic_direction_statement_v2(
    snapshot,
) -> str | None:
    if snapshot.status != PeriodicMetricStatusV2.READY:
        return None

    name = snapshot.spec.chinese_name

    if (
        snapshot.spec.display_kind
        == PeriodicMetricDisplayKindV2.RATIO
        and snapshot.percentage_point_change is not None
    ):
        value = snapshot.percentage_point_change

        if value > 0:
            direction = "提升"
        elif value < 0:
            direction = "下降"
        else:
            direction = "持平"

        if value == 0:
            return f"{name}较参考期持平。"

        return (
            f"{name}较参考期{direction}"
            f"{abs(value):.2f} 个百分点。"
        )

    change = snapshot.relative_change

    if change is None:
        return (
            f"{name}当前值为"
            f"{_periodic_metric_value_text(snapshot, reference=False)}。"
        )

    if change > 0:
        direction = "上升"
    elif change < 0:
        direction = "下降"
    else:
        direction = "持平"

    if change == 0:
        return f"{name}较参考期持平。"

    return (
        f"{name}较参考期{direction}"
        f"{abs(change) * Decimal('100'):.2f}%。"
    )


def _periodic_summary_lines_v2(
    report: PeriodicBusinessReportV2,
) -> tuple[str, ...]:
    """
    Deterministic business observations from already-computed deltas.

    不重新计算 KPI，不声明业务因果。
    """

    lookup = _periodic_metric_lookup_v2(
        report
    )
    lines: list[str] = []

    for metric_name in (
        "gmv",
        "buyer_count",
        "order_count",
        "spending_per_buyer",
        "aus",
        "r12_repurchase_rate",
    ):
        snapshot = lookup.get(
            metric_name
        )
        if snapshot is None:
            continue

        statement = _periodic_direction_statement_v2(
            snapshot
        )
        if statement is not None:
            lines.append(statement)

    failed = tuple(
        item.spec.chinese_name
        for item in report.metrics
        if item.status != PeriodicMetricStatusV2.READY
    )

    if failed:
        lines.append(
            "当前不可释放指标："
            + "、".join(failed)
            + "；已在“本期限制”中单独披露。"
        )

    return tuple(lines)


def _add_periodic_executive_summary_v2(
    document: Document,
    report: PeriodicBusinessReportV2,
) -> None:
    lookup = _periodic_metric_lookup_v2(
        report
    )

    priority = tuple(
        lookup[name]
        for name in (
            "gmv",
            "buyer_count",
            "order_count",
            "r12_repurchase_rate",
        )
        if (
            name in lookup
            and lookup[name].status
            == PeriodicMetricStatusV2.READY
        )
    )

    if priority:
        table = document.add_table(
            rows=1,
            cols=4,
        )
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = (
            "指标",
            "当前期",
            "变化",
            "变化率",
        )
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        _set_table_header(table.rows[0])

        for snapshot in priority:
            row = table.add_row().cells
            row[0].text = snapshot.spec.chinese_name
            row[1].text = _periodic_metric_value_text(
                snapshot,
                reference=False,
            )
            row[2].text = _periodic_change_text(
                snapshot
            )
            row[3].text = (
                _ratio_text(snapshot.relative_change)
                if snapshot.relative_change is not None
                else "—"
            )

    document.add_heading(
        "经营观察",
        level=2,
    )

    _add_bullets(
        document,
        list(
            _periodic_summary_lines_v2(
                report
            )
        ),
    )

def render_investigation_business_docx_v2(
    report: InvestigationReportV2,
) -> bytes:
    """
    Mode-aware business-facing DOCX.

    Investigation 采用：
    决策摘要 -> 核心对比 -> 调查明细 -> 边界 / 下一步。
    """

    document = Document()
    _set_document_defaults(document)

    title = _business_report_title_v2(
        report
    )
    document.core_properties.title = title

    _add_title(
        document,
        title,
        "基于受治理数据与结构化 Evidence 的确定性交付",
    )

    section_index = 0

    def add_section(
        title_text: str,
    ) -> None:
        nonlocal section_index
        section_index += 1

        numbers = {
            1: "一",
            2: "二",
            3: "三",
            4: "四",
            5: "五",
            6: "六",
            7: "七",
            8: "八",
            9: "九",
            10: "十",
        }

        document.add_heading(
            f"{numbers[section_index]}、{title_text}",
            level=1,
        )

    add_section("本次业务问题")
    document.add_paragraph(
        report.original_question
    )

    if report.resolution_note:
        document.add_paragraph(
            f"补充口径：{report.resolution_note}"
        )

    mode_group = _report_mode_group_v2(
        report
    )

    if mode_group == "fact":
        add_section("核心结果")
        document.add_paragraph(
            report.answer_snapshot
        )

        if _public_fact_compositions_v2(report):
            add_section("构成与验证")
            _add_fact_composition_sections_v2(
                document,
                report,
            )

        add_section("业务口径与可信边界")
        _add_business_trust_section_v2(
            document,
            report,
        )

    elif mode_group == "comparison":
        if report.comparison_summary is not None:
            add_section("核心对比")
            _add_comparison_section(
                document,
                report,
            )
        else:
            add_section("核心结果")
            document.add_paragraph(
                report.answer_snapshot
            )

        findings = tuple(
            item.summary
            for item in (
                report.executive_brief.key_findings
            )
        )

        if findings:
            add_section("关键发现")
            _add_bullets(
                document,
                list(findings),
            )

        add_section("业务口径与可信边界")
        _add_business_trust_section_v2(
            document,
            report,
        )

    else:
        add_section("决策摘要")
        _add_investigation_decision_summary_v2(
            document,
            report,
        )

        if report.comparison_summary is not None:
            add_section("核心对比")
            _add_comparison_section(
                document,
                report,
            )
        else:
            add_section("核心结果")
            document.add_paragraph(
                report.answer_snapshot
            )

        all_steps = _investigation_all_steps_v2(
            report
        )

        if all_steps:
            add_section("调查明细")

            for index, step in enumerate(
                report.investigation_steps,
                start=1,
            ):
                _add_change_step(
                    document,
                    step=step,
                    sequence=index,
                    source_label="受控调查",
                )

            offset = len(
                report.investigation_steps
            )

            for index, step in enumerate(
                report.user_exploration_steps,
                start=1,
            ):
                _add_change_step(
                    document,
                    step=step,
                    sequence=offset + index,
                    source_label="用户主动探索",
                )

        confirmed = list(
            _investigation_confirmed_v2(
                report
            )
        )
        cannot_confirm = list(
            _investigation_cannot_confirm_v2(
                report
            )
        )
        next_steps = list(
            _investigation_next_steps_v2(
                report
            )
        )

        if confirmed:
            add_section("可以确认")
            _add_bullets(
                document,
                confirmed,
            )

        if cannot_confirm:
            add_section("暂不能确认与因果边界")
            _add_bullets(
                document,
                cannot_confirm,
            )

        if next_steps:
            add_section("下一步建议")
            _add_bullets(
                document,
                next_steps,
            )

        add_section("业务口径与可信边界")
        _add_business_trust_section_v2(
            document,
            report,
        )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()




def _cadence_label(
    report: PeriodicBusinessReportV2,
) -> str:
    return {
        "daily": "日报",
        "weekly": "周报",
        "monthly": "月报",
    }.get(
        report.cadence.value,
        report.cadence.value,
    )


def _add_periodic_metric_table(
    document: Document,
    *,
    title: str,
    metrics,
) -> None:
    document.add_heading(
        title,
        level=2,
    )

    table = document.add_table(
        rows=1,
        cols=5,
    )
    table.style = "Table Grid"
    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    headers = (
        "指标",
        "参考期",
        "当前期",
        "变化",
        "变化率",
    )
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    _set_table_header(table.rows[0])

    for snapshot in metrics:
        row = table.add_row().cells
        row[0].text = snapshot.spec.chinese_name
        row[1].text = _periodic_metric_value_text(
            snapshot,
            reference=True,
        )
        row[2].text = _periodic_metric_value_text(
            snapshot,
            reference=False,
        )
        row[3].text = _periodic_change_text(snapshot)
        row[4].text = (
            _ratio_text(snapshot.relative_change)
            if snapshot.status
            == PeriodicMetricStatusV2.READY
            else "—"
        )

def render_periodic_business_docx_v2(
    report: PeriodicBusinessReportV2,
) -> bytes:
    """
    PeriodicBusinessReportV2 -> business-facing DOCX.

    新增 deterministic “本期经营摘要”，
    只消费现有 Snapshot 的 current / delta / rate。
    """

    document = Document()
    _set_document_defaults(document)

    cadence_label = _cadence_label(report)
    document.core_properties.title = (
        f"AI Data Analyst {cadence_label}"
    )

    _add_title(
        document,
        f"AI Data Analyst · {cadence_label}",
        f"锚点日期：{report.anchor_date}",
    )

    document.add_heading(
        "一、报表范围",
        level=1,
    )
    document.add_paragraph(
        "参考期："
        f"{report.comparison.reference_window.start_date} "
        "至 "
        f"{report.comparison.reference_window.end_date}"
    )
    document.add_paragraph(
        "当前期："
        f"{report.comparison.current_window.start_date} "
        "至 "
        f"{report.comparison.current_window.end_date}"
    )

    document.add_heading(
        "二、本期经营摘要",
        level=1,
    )
    _add_periodic_executive_summary_v2(
        document,
        report,
    )

    section_titles = {
        PeriodicMetricSectionV2.OVERVIEW: "三、经营概览",
        PeriodicMetricSectionV2.SALES_DRIVER: "四、销售驱动",
        PeriodicMetricSectionV2.CUSTOMER_HEALTH: "五、客户健康",
    }

    for section in (
        PeriodicMetricSectionV2.OVERVIEW,
        PeriodicMetricSectionV2.SALES_DRIVER,
        PeriodicMetricSectionV2.CUSTOMER_HEALTH,
    ):
        metrics = tuple(
            item
            for item in report.metrics
            if item.spec.section == section
        )

        if metrics:
            _add_periodic_metric_table(
                document,
                title=section_titles[section],
                metrics=metrics,
            )

    document.add_heading(
        "六、驱动关系验证",
        level=1,
    )

    reconciliation_labels = {
        "reconciled": "已对账",
        "not_reconciled": "未完全对账",
        "not_applicable": "当前不可验证",
    }

    if report.driver_reconciliations:
        for item in report.driver_reconciliations:
            raw_status = (
                item.status.value
                if hasattr(item.status, "value")
                else str(item.status)
            )
            document.add_paragraph(
                f"{item.relationship}｜"
                f"{reconciliation_labels.get(raw_status, raw_status)}"
            )
    else:
        document.add_paragraph(
            "当前没有可释放的驱动关系验证。"
        )

    failed = tuple(
        item
        for item in report.metrics
        if item.status != PeriodicMetricStatusV2.READY
    )

    document.add_heading(
        "七、本期限制",
        level=1,
    )

    if failed:
        for item in failed:
            document.add_paragraph(
                f"{item.spec.chinese_name}："
                "本期未形成可安全释放的结果。",
                style="List Bullet",
            )
    else:
        document.add_paragraph(
            "当前已注册指标均可正常展示。"
        )

    document.add_heading(
        "八、指标说明",
        level=1,
    )

    for item in report.metrics:
        document.add_paragraph(
            f"{item.spec.chinese_name}："
            f"{item.spec.purpose}",
            style="List Bullet",
        )

    document.add_paragraph(
        "说明：本期经营摘要只描述结构化指标已经确认的变化方向；"
        "驱动关系验证是确定性算术核对，均不代表因果解释。"
    )

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
